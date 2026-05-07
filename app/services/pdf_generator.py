import io
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from app.models.schemas import ClientInfo

CHECKED_NUM_ID = "1"
UNCHECKED_NUM_ID = "3"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

DECLARE_ITEMS = [
    (
        "dichiara_norma",
        "seguito la norma tecnica applicabile",
    ),
    (
        "dichiara_componenti",
        "installato componenti e materiali adatti",
    ),
    (
        "dichiara_controllo",
        "controllato l'impianto ai fini",
    ),
]

ALLEGATI_ITEMS = [
    ("allegato_progetto", "progetto ai sensi"),
    ("allegato_relazione", "relazione con tipologie"),
    ("allegato_schema", "schema di impianto"),
    ("allegato_precedenti", "riferimento a dichiarazione"),
    ("allegato_certificato", "copia del certificato"),
    ("allegato_conformita", "attestazione di conformita"),
]


class PDFTemplateError(Exception):
    pass


@dataclass(frozen=True)
class DeclarationDocxData:
    esecutrice_di: str
    inteso_come: str
    commissionato_da: str
    comune: str
    via: str
    proprietario: str
    adibito_ad_uso: str
    data: str
    allegati: dict[str, bool]


def get_template_fields(template_path: Path) -> list[str]:
    """Return DOCX placeholder names found in the template."""
    if not template_path.exists():
        raise PDFTemplateError(f"Template DOCX non trovato: {template_path}")

    document = Document(str(template_path))
    text = "\n".join(paragraph.text for paragraph in _iter_paragraphs(document))
    fields: list[str] = []
    for key in (
        "esecutrice_di",
        "inteso_come",
        "commissionato_da",
        "comune",
        "via",
        "adibito_ad_uso",
    ):
        if "{{" + key + "}}" in text:
            fields.append(key)
    return fields


def render_declaration_docx(
    client: ClientInfo,
    template_path: Path,
    extra_fields: dict[str, str] | None = None,
    allegati: dict[str, bool] | None = None,
) -> bytes:
    """Fill the DOCX template as a real Word document and return DOCX bytes."""
    if not template_path.exists():
        raise PDFTemplateError(f"Template DOCX non trovato: {template_path}")

    data = _build_docx_data(client, extra_fields, allegati)
    document = Document(str(template_path))

    _replace_placeholders(document, data)
    _apply_tick_markers(document, data.allegati)
    _remove_blank_line_after_allegati(document)
    _replace_footer_date_and_signature(document, data.data)
    _insert_signature_images(document, template_path)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def generate_declaration(
    client: ClientInfo,
    template_path: Path,
    extra_fields: dict[str, str] | None = None,
    allegati: dict[str, bool] | None = None,
) -> bytes:
    """Fill the Word template and export a one-page PDF."""
    docx_bytes = render_declaration_docx(client, template_path, extra_fields, allegati)

    with tempfile.TemporaryDirectory(prefix="dichiarazione_") as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / "dichiarazione.docx"
        docx_path.write_bytes(docx_bytes)
        return _convert_docx_to_pdf(docx_path, tmp_path)


def _build_docx_data(
    client: ClientInfo,
    extra_fields: dict[str, str] | None,
    allegati: dict[str, bool] | None,
) -> DeclarationDocxData:
    extra_fields = extra_fields or {}
    allegati = allegati or {}

    proprietario = extra_fields.get("proprietario") or client.name
    all_tick_names = [name for name, _ in DECLARE_ITEMS + ALLEGATI_ITEMS]
    always_checked = {name for name, _ in DECLARE_ITEMS}

    return DeclarationDocxData(
        esecutrice_di=extra_fields.get("tipo_impianto") or "",
        inteso_come=extra_fields.get("descrizione_impianto") or "",
        commissionato_da=client.name,
        comune=extra_fields.get("comune_installazione") or client.address_city or "",
        via=extra_fields.get("via_installazione") or client.address_street or "",
        proprietario=proprietario,
        adibito_ad_uso=extra_fields.get("uso_edificio") or "",
        data=extra_fields.get("data") or "",
        allegati={
            name: True if name in always_checked else bool(allegati.get(name, False))
            for name in all_tick_names
        },
    )


def _iter_paragraphs(document: DocxDocument):
    for paragraph in document.paragraphs:
        yield paragraph
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _replace_placeholders(document: DocxDocument, data: DeclarationDocxData) -> None:
    replacements = {
        "{{esecutrice_di}}": data.esecutrice_di,
        "{{inteso_come}}": data.inteso_come,
        "{{commissionato_da}}": data.commissionato_da,
        "{{comune}}": data.comune,
        "{{via}}": data.via,
        "{{adibito_ad_uso}}": data.adibito_ad_uso,
    }

    for paragraph in _iter_paragraphs(document):
        for run in paragraph.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    _replace_placeholder_in_run(run, placeholder, value)

        if data.proprietario and data.proprietario != data.commissionato_da:
            _replace_phrase_with_bold_value(
                paragraph,
                "di proprieta dello stesso",
                "di proprieta di ",
                data.proprietario,
            )
            _replace_phrase_with_bold_value(
                paragraph,
                "di proprietà dello stesso",
                "di proprietà di ",
                data.proprietario,
            )


def _apply_tick_markers(document: DocxDocument, allegati: dict[str, bool]) -> None:
    for field_name, text_marker in DECLARE_ITEMS + ALLEGATI_ITEMS:
        paragraph = _find_paragraph(document, text_marker)
        if paragraph is not None:
            _set_numbering_marker(paragraph, checked=allegati.get(field_name, False))


def _replace_footer_date_and_signature(document: DocxDocument, data: str) -> None:
    if not data:
        return

    for paragraph in _iter_paragraphs(document):
        if "Data " in paragraph.text and "Il responsabile tecnico" in paragraph.text:
            match = re.search(r"\d{2}/\d{2}/\d{4}", paragraph.text)
            if match:
                _replace_phrase_with_bold_value(paragraph, match.group(0), "", data)
            return


def _insert_signature_images(document: DocxDocument, template_path: Path) -> None:
    signature_path = template_path.parent / "firma.png"
    if not signature_path.exists():
        raise PDFTemplateError(f"Immagine firma non trovata: {signature_path}")

    for paragraph in _iter_paragraphs(document):
        if "Data " in paragraph.text and "Il responsabile tecnico" in paragraph.text:
            _add_floating_picture(
                paragraph,
                signature_path,
                width=Inches(1.25),
                x_offset=Inches(1.78),
                y_offset=Inches(0.03),
            )
            _add_floating_picture(
                paragraph,
                signature_path,
                width=Inches(1.25),
                x_offset=Inches(4.05),
                y_offset=Inches(0.03),
            )
            return


def _add_floating_picture(paragraph, image_path: Path, width, x_offset, y_offset) -> None:
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)
    inline = run._r.find(".//" + qn("wp:inline"))
    if inline is None:
        raise PDFTemplateError("Impossibile creare la firma flottante nel DOCX")
    inline.getparent().replace(
        inline,
        _floating_anchor_from_inline(inline, x_offset=x_offset, y_offset=y_offset),
    )


def _floating_anchor_from_inline(inline, x_offset, y_offset):
    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0")
    anchor.set("distB", "0")
    anchor.set("distL", "0")
    anchor.set("distR", "0")
    anchor.set("simplePos", "0")
    anchor.set("relativeHeight", "251659264")
    anchor.set("behindDoc", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "column")
    pos_h_offset = OxmlElement("wp:posOffset")
    pos_h_offset.text = str(int(x_offset))
    position_h.append(pos_h_offset)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    pos_v_offset = OxmlElement("wp:posOffset")
    pos_v_offset.text = str(int(y_offset))
    position_v.append(pos_v_offset)
    anchor.append(position_v)

    for tag in ("wp:extent",):
        element = inline.find(qn(tag))
        if element is not None:
            anchor.append(deepcopy(element))

    effect_extent = OxmlElement("wp:effectExtent")
    effect_extent.set("l", "0")
    effect_extent.set("t", "0")
    effect_extent.set("r", "0")
    effect_extent.set("b", "0")
    anchor.append(effect_extent)

    anchor.append(OxmlElement("wp:wrapNone"))

    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        element = inline.find(qn(tag))
        if element is not None:
            anchor.append(deepcopy(element))

    return anchor


def _find_paragraph(document: DocxDocument, text_marker: str):
    normalized_marker = _normalize(text_marker)
    for paragraph in _iter_paragraphs(document):
        if normalized_marker in _normalize(paragraph.text):
            return paragraph
    return None


def _normalize(value: str) -> str:
    return (
        value.replace("à", "a")
        .replace("è", "e")
        .replace("é", "e")
        .replace("ì", "i")
        .replace("ò", "o")
        .replace("ù", "u")
        .replace("’", "'")
        .lower()
    )


def _set_numbering_marker(paragraph, checked: bool) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")

    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        num_id = OxmlElement("w:numId")
        num_pr.append(num_id)
    num_id.set(qn("w:val"), CHECKED_NUM_ID if checked else UNCHECKED_NUM_ID)


def _replace_placeholder_in_run(run, placeholder: str, value: str) -> None:
    before, after = run.text.split(placeholder, 1)
    run.text = before

    value_r = _make_text_run(value, run, bold=True)
    run._r.addnext(value_r)

    if after:
        after_r = _make_text_run(after, run)
        value_r.addnext(after_r)


def _replace_phrase_with_bold_value(paragraph, old: str, prefix: str, value: str) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    start = full_text.find(old)
    if start == -1:
        return False

    end = start + len(old)
    cursor = 0
    inserted_value = False

    for run in paragraph.runs:
        text = run.text
        run_start = cursor
        run_end = cursor + len(text)
        cursor = run_end

        if run_end <= start or run_start >= end:
            continue

        before = text[: max(0, start - run_start)] if run_start <= start else ""
        after = text[max(0, end - run_start) :] if run_end >= end else ""

        if not inserted_value:
            run.text = before + prefix
            value_r = _make_text_run(value, run, bold=True)
            run._r.addnext(value_r)
            if after:
                value_r.addnext(_make_text_run(after, run))
            inserted_value = True
        else:
            run.text = after

    return True


def _make_text_run(text: str, source_run, bold: bool | None = None):
    r = OxmlElement("w:r")
    if source_run._r.rPr is not None:
        r_pr = deepcopy(source_run._r.rPr)
    else:
        r_pr = OxmlElement("w:rPr")

    if bold is True:
        b = r_pr.find(qn("w:b"))
        if b is None:
            b = OxmlElement("w:b")
            r_pr.append(b)
        b.set(qn("w:val"), "1")

    r.append(r_pr)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    return r


def _set_paragraph_text(paragraph, text: str) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    paragraph.add_run(text)


def _replace_whole_paragraph_text(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    _set_paragraph_text(paragraph, paragraph.text.replace(old, new))


def _replace_text_across_runs(paragraph, old: str, new: str) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    start = full_text.find(old)
    if start == -1:
        return False

    end = start + len(old)
    cursor = 0
    wrote_replacement = False

    for run in paragraph.runs:
        text = run.text
        run_start = cursor
        run_end = cursor + len(text)
        cursor = run_end

        if run_end <= start or run_start >= end:
            continue

        prefix = text[: max(0, start - run_start)] if run_start <= start else ""
        suffix = text[max(0, end - run_start) :] if run_end >= end else ""
        if not wrote_replacement:
            run.text = prefix + new + suffix
            wrote_replacement = True
        else:
            run.text = suffix

    return True


def _remove_blank_line_after_allegati(document: DocxDocument) -> None:
    paragraphs = list(document.paragraphs)
    for index, paragraph in enumerate(paragraphs[:-1]):
        if "Allegati obbligatori:" in paragraph.text and not paragraphs[index + 1].text.strip():
            _delete_paragraph(paragraphs[index + 1])
            return


def _remove_blank_lines_around_declina(document: DocxDocument) -> None:
    paragraphs = list(document.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != "DECLINA":
            continue
        if index > 0 and not paragraphs[index - 1].text.strip():
            _delete_paragraph(paragraphs[index - 1])
        if index + 1 < len(paragraphs) and not paragraphs[index + 1].text.strip():
            _delete_paragraph(paragraphs[index + 1])
        return


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def _convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> bytes:
    converters = [
        _convert_with_libreoffice,
        _convert_with_pages,
    ]
    errors: list[str] = []

    for converter in converters:
        try:
            pdf_path = converter(docx_path, output_dir)
        except FileNotFoundError as exc:
            errors.append(str(exc))
            continue
        except subprocess.SubprocessError as exc:
            errors.append(str(exc))
            continue

        if pdf_path.exists():
            return pdf_path.read_bytes()
        errors.append(f"Conversione non riuscita: {pdf_path} non trovato")

    raise PDFTemplateError(
        "Impossibile esportare il DOCX in PDF. Installa LibreOffice oppure usa Pages.app. "
        + " | ".join(errors)
    )


def _convert_with_libreoffice(docx_path: Path, output_dir: Path) -> Path:
    soffice = (
        shutil.which("soffice")
        or shutil.which("libreoffice")
        or _existing_path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    )
    if not soffice:
        raise FileNotFoundError("LibreOffice/soffice non trovato")

    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        timeout=60,
    )
    return output_dir / (docx_path.stem + ".pdf")


def _convert_with_pages(docx_path: Path, output_dir: Path) -> Path:
    pages_app = Path("/Applications/Pages.app")
    if not pages_app.exists() or not shutil.which("osascript"):
        raise FileNotFoundError("Pages.app/osascript non trovato")

    pdf_path = output_dir / (docx_path.stem + ".pdf")
    script = f"""
tell application "Pages"
    set theDoc to open POSIX file "{docx_path}"
    export theDoc to POSIX file "{pdf_path}" as PDF
    close theDoc saving no
end tell
"""
    subprocess.run(
        ["osascript", "-e", script],
        check=True,
        capture_output=True,
        timeout=60,
    )
    return pdf_path


def _existing_path(path: str) -> str | None:
    candidate = Path(path)
    return str(candidate) if candidate.exists() else None
