import io
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.services import pdf_generator
from app.services.pdf_generator import (
    CHECKED_NUM_ID,
    PDFTemplateError,
    UNCHECKED_NUM_ID,
    generate_declaration,
    get_template_fields,
    render_declaration_docx,
)


def _docx_from_bytes(docx_bytes: bytes):
    return Document(io.BytesIO(docx_bytes))


def _text(document) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _paragraph_num_id(document, text_marker: str) -> str | None:
    for paragraph in document.paragraphs:
        if text_marker in paragraph.text:
            num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
            if num_pr is None:
                return None
            num_id = num_pr.find(qn("w:numId"))
            return num_id.get(qn("w:val")) if num_id is not None else None
    return None


def _run_is_bold(document, exact_text: str) -> bool:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.text == exact_text:
                return bool(run.bold)
    return False


def test_get_template_fields(sample_template_docx: Path):
    fields = get_template_fields(sample_template_docx)
    assert "commissionato_da" in fields
    assert "comune" in fields
    assert "via" in fields
    assert "inteso_come" in fields


def test_missing_template(tmp_path: Path):
    with pytest.raises(PDFTemplateError):
        get_template_fields(tmp_path / "nope.docx")


def test_docx_placeholders_are_replaced(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(
        sample_client,
        sample_template_docx,
        extra_fields={
            "data": "15/04/2026",
            "tipo_impianto": "nuovo impianto idraulico",
            "descrizione_impianto": "impianto idrico sanitario",
            "uso_edificio": "civile",
        },
    )
    document = _docx_from_bytes(docx_bytes)
    text = _text(document)

    assert "{{" not in text
    assert "Mario Rossi SRL" in text
    assert "Roma" in text
    assert "Via Roma 1" in text
    assert "nuovo impianto idraulico" in text
    assert "impianto idrico sanitario" in text
    assert "civile" in text
    assert "Data 15/04/2026" in text
    assert "Il responsabile tecnico" in text
    assert "Il dichiarante" in text


def test_app_inserted_values_are_bold(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(
        sample_client,
        sample_template_docx,
        extra_fields={
            "data": "15/04/2026",
            "tipo_impianto": "nuovo impianto idraulico",
            "descrizione_impianto": "impianto idrico sanitario",
            "uso_edificio": "civile",
        },
    )
    document = _docx_from_bytes(docx_bytes)

    assert _run_is_bold(document, "Mario Rossi SRL")
    assert _run_is_bold(document, "Roma")
    assert _run_is_bold(document, "Via Roma 1")
    assert _run_is_bold(document, "nuovo impianto idraulico")
    assert _run_is_bold(document, "impianto idrico sanitario")
    assert _run_is_bold(document, "civile")
    assert _run_is_bold(document, "15/04/2026")


def test_proprietario_defaults_to_dello_stesso(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(sample_client, sample_template_docx)
    assert "di proprietà dello stesso" in _text(_docx_from_bytes(docx_bytes))


def test_proprietario_override(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(
        sample_client,
        sample_template_docx,
        extra_fields={"proprietario": "Sig. Bianchi"},
    )
    document = _docx_from_bytes(docx_bytes)
    assert "di proprietà di Sig. Bianchi" in _text(document)
    assert _run_is_bold(document, "Sig. Bianchi")


def test_installation_address_override(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(
        sample_client,
        sample_template_docx,
        extra_fields={"comune_installazione": "Voghera", "via_installazione": "via Leopardi 11"},
    )
    text = _text(_docx_from_bytes(docx_bytes))
    assert "Voghera" in text
    assert "via Leopardi 11" in text


def test_tick_markers_are_word_numbering(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(
        sample_client,
        sample_template_docx,
        allegati={
            "dichiara_norma": True,
            "dichiara_componenti": False,
            "dichiara_controllo": True,
            "allegato_progetto": True,
            "allegato_relazione": False,
            "allegato_schema": True,
            "allegato_precedenti": False,
            "allegato_certificato": True,
            "allegato_conformita": False,
        },
    )
    document = _docx_from_bytes(docx_bytes)

    assert _paragraph_num_id(document, "seguito la norma tecnica") == CHECKED_NUM_ID
    assert _paragraph_num_id(document, "installato componenti") == CHECKED_NUM_ID
    assert _paragraph_num_id(document, "controllato l") == CHECKED_NUM_ID
    assert _paragraph_num_id(document, "copia del certificato") == CHECKED_NUM_ID
    assert _paragraph_num_id(document, "attestazione di conformità") == UNCHECKED_NUM_ID


def test_required_attachments_default_to_unchecked(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(sample_client, sample_template_docx)
    document = _docx_from_bytes(docx_bytes)

    assert _paragraph_num_id(document, "seguito la norma tecnica") == CHECKED_NUM_ID
    assert _paragraph_num_id(document, "progetto ai sensi") == UNCHECKED_NUM_ID
    assert _paragraph_num_id(document, "relazione con tipologie") == UNCHECKED_NUM_ID
    assert _paragraph_num_id(document, "schema di impianto") == UNCHECKED_NUM_ID
    assert _paragraph_num_id(document, "copia del certificato") == UNCHECKED_NUM_ID


def test_blank_line_below_allegati_heading_is_removed(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(sample_client, sample_template_docx)
    document = _docx_from_bytes(docx_bytes)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    allegati_index = paragraphs.index("Allegati obbligatori:")

    assert paragraphs[allegati_index + 1].strip()


def test_signature_is_inserted_for_both_signers(sample_template_docx: Path, sample_client):
    docx_bytes = render_declaration_docx(sample_client, sample_template_docx)
    document = _docx_from_bytes(docx_bytes)
    anchors = document.element.findall(".//" + qn("wp:anchor"))

    assert len(anchors) == 2
    assert len(document.inline_shapes) == 0


def test_generate_declaration_exports_pdf(sample_template_docx: Path, sample_client, monkeypatch):
    captured = {}

    def fake_converter(docx_path: Path, output_dir: Path) -> bytes:
        captured["docx_path"] = docx_path
        captured["output_dir"] = output_dir
        generated = Document(str(docx_path))
        assert "Mario Rossi SRL" in _text(generated)
        return b"%PDF-1.4\n%%EOF"

    monkeypatch.setattr(pdf_generator, "_convert_docx_to_pdf", fake_converter)

    pdf_bytes = generate_declaration(sample_client, sample_template_docx)

    assert pdf_bytes.startswith(b"%PDF")
    assert captured["docx_path"].name == "dichiarazione.docx"


def test_missing_template_generate(tmp_path: Path, sample_client):
    with pytest.raises(PDFTemplateError):
        render_declaration_docx(sample_client, tmp_path / "missing.docx")
